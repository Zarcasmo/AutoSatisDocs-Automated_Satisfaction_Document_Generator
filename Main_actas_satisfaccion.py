# -*- coding: utf-8 -*-
"""
Created on Tue Feb 25 10:44:13 2025

@author: alopagui

Script para la conversión de archivos Word a PDF de manera automatizada.

Este script toma un conjunto de archivos .docx y los convierte en archivos .pdf utilizando la automatización de Microsoft Word a través de win32com.client. 
Incluye una barra de progreso para monitorear el avance del proceso y muestra el tiempo transcurrido.

Requisitos:
- Python 3
- Bibliotecas: os, time, colorama, win32com.client

Uso:
1. Asegurar que Microsoft Word está instalado en el sistema.
2. Ubicar los archivos .docx en la carpeta de origen.
3. Ejecutar el script para convertir los archivos en formato PDF.
"""

import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import comtypes.client  
import time
import sys
from colorama import Fore, Style, init

# Inicializar colorama para colores en la consola
init(autoreset=True)

#Tamaño que tendran las firmas
tam_imagen=1

#Guardar documento .docx
borrar_docx = True

# Cargar el archivo Excel con los datos de entrada
# Este archivo debe contener las columnas necesarias como "Nombre", "Cedula", "Calidad", "Municipio", etc.
df = pd.read_excel("Intervencion_forestal_guaduales_SHP.xlsx")
df = df.fillna('')

#Información de las firmas de los lideres de grupos
df_lideres = pd.read_excel("Firmas_lideres.xlsx")

#Merge para obtener la ruta de las firmas de los lideres
df = df.merge(df_lideres[['NOMBRE_LIDER', 'FIRMA_LIDER']], left_on='jefeCuadrilla', right_on='NOMBRE_LIDER', how='left')
# Eliminar la columna duplicada si no la necesitas
df.drop(columns=['NOMBRE_LIDER'], inplace=True)


# Ruta de la plantilla Word que se usará como base para generar los documentos
template_path = "formato_socializa.docx"
output_folder = "output_pdfs" # Carpeta donde se guardarán los archivos generados
firmas_folder = "Firmas"  # Carpeta donde están las imágenes de firmas
firmas_lideres_folder = "Firmas lideres grupos"  # Carpeta donde están las imágenes de firmas de los lideres
os.makedirs(output_folder, exist_ok=True) # Crear la carpeta de salida si no existe

# Archivo Excel de salida con la consolidación de resultados
output_excel_path = os.path.join(output_folder, "Resumen_Resultados.xlsx")

# Inicializar lista para almacenar los resultados
resultados = []

# Función para imprimir una barra de progreso durante la conversión a PDF
def print_progress_bar(iteration, total, length=50, start_time=None):
    """
    Imprime una barra de progreso en la consola y muestra el tiempo transcurrido al completar.

    :param iteration: El número actual de iteraciones.
    :param total: El número total de iteraciones.
    :param length: La longitud de la barra de progreso.
    :param start_time: El tiempo de inicio de la ejecución.
    """
    progress = (iteration / total)
    arrow = '█' * int(round(progress * length))
    spaces = '░' * (length - len(arrow))
    percent = int(round(progress * 100))
    
    # Construcción de la barra con colores
    bar = f'[{Fore.GREEN}{arrow}{spaces}{Style.RESET_ALL}]'
    percent_text = f'{Fore.CYAN}{percent}%{Style.RESET_ALL}'
    
    if start_time:
        elapsed_time = time.time() - start_time
        elapsed_time_formatted = time.strftime("%H:%M:%S", time.gmtime(elapsed_time))
        sys.stdout.write(f'\r{bar} {percent_text} Complete - Time elapsed: {elapsed_time_formatted}')
    else:
        sys.stdout.write(f'\r{bar} {percent_text} Complete')
    
    # Si ha llegado al 100%, imprime una nueva línea
    if iteration == total:
        sys.stdout.write('\n')
    
    sys.stdout.flush()

# Función mejorada para reemplazar texto manteniendo formato
def replace_text_keep_format(doc, old_text, new_text):
    """
    Reemplaza texto específico manteniendo el formato original y el resto del contenido.
    
    :param doc: Documento de Word (objeto Document)
    :param old_text: Texto a reemplazar
    :param new_text: Texto de reemplazo
    """
    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, str(new_text))

    # Reemplazar en tablas de manera recursiva
    def replace_in_table(table):
        for row in table.rows:
            for cell in row.cells:
                # Reemplazar en párrafos de la celda
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if old_text in run.text:
                            run.text = run.text.replace(old_text, str(new_text))
                
                # Reemplazar en subtablas anidadas si existen
                nested_tables = cell.tables
                for nested_table in nested_tables:
                    replace_in_table(nested_table)

    # Buscar y reemplazar en todas las tablas del documento
    for table in doc.tables:
        replace_in_table(table)

# Función para insertar imágenes en el documento Word reemplazando texto específico
def replace_text_with_image(doc, placeholder, image_path, width=tam_imagen):
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(image_path, width=Inches(width))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    cell.text = ""  
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(image_path, width=Inches(width))

# Listas para almacenar rutas de archivos generados
word_files = []
pdf_files = []
cuenta_word = 0
total = len(df)-1
start_time = time.time()  # Captura el tiempo de inicio
print("Proceso de generación de Words")

# Procesar cada fila del Excel
for index, row in df.iterrows():
    word_output_path = os.path.join(output_folder, f"documento_{index}.docx")
    pdf_output_path = os.path.join(output_folder, f"documento_{index}.pdf")
    status = "Éxito"
    error_msg = ""

    try:
        # Crear un nuevo documento basado en la plantilla
        doc = Document(template_path)

        #Datos usuarios
        replace_text_keep_format(doc, "@NOMBRE@", row["nombreUsuario"])
        replace_text_keep_format(doc, "@CEDULA@", str(row["cedulaUsuario"]))
        replace_text_keep_format(doc, "@CALIDAD@", row["calidadFirma"])
        #Ubicacion
        if row["MUNICIPIO_FINCA"] != '':
            replace_text_keep_format(doc, "@MUNICIPIO@", str(row["MUNICIPIO_FINCA"]))
        else:
            replace_text_keep_format(doc, "@MUNICIPIO@", str(row["Municipio_manual"]))
        replace_text_keep_format(doc, "@VEREDA@", row["Vereda"])
        replace_text_keep_format(doc, "@DIRECCION@", row["Finca_Direccion"])
        replace_text_keep_format(doc, "@LATITUD@", str(row["COORDENADAS_FINCA"].split(',')[0]))
        replace_text_keep_format(doc, "@LONGITUD@", str(row["COORDENADAS_FINCA"].split(',')[1]))
        #Comentarios
        replace_text_keep_format(doc, "@COMENTARIO_EDEQ@", row["Observacion"])
        replace_text_keep_format(doc, "@COMENTARIO_USUARIO@", row["observacionUsuario"])
        #Fecha
        replace_text_keep_format(doc, "@YEAR@", str(row["Fecha_accion"].year))
        replace_text_keep_format(doc, "@MONTH@", str(row["Fecha_accion"].month))
        replace_text_keep_format(doc, "@DAY@", str(row["Fecha_accion"].day))
        #OT/Evento
        if row["OTMX"] != '':
            replace_text_keep_format(doc, "@OT/EVENTO@", str(int(row["OTMX"])))
        else:
            replace_text_keep_format(doc, "@OT/EVENTO@", str(row["Evento_SP7"]))
        #Cantidades
        replace_text_keep_format(doc, "@PODA@", str(row["cantidadPodasFirmadas"]))
        replace_text_keep_format(doc, "@RETIROS@", str(row["cantidadRetirosFirmados"]))
        replace_text_keep_format(doc, "@GUADUA@", str(row["cantidadRenuevosFirmados"]))
        replace_text_keep_format(doc, "@ROCERIA@", str(row["metrosRoceriaFirmados"]))
        #Residuos
        if row["disposicionResiduos"]=='A cargo de EDEQ':
            replace_text_keep_format(doc, "@Resi_EDEQ@", "X")
            replace_text_keep_format(doc, "@Resi_user@", "")
        else:
            replace_text_keep_format(doc, "@Resi_EDEQ@", "")
            replace_text_keep_format(doc, "@Resi_user@", "X")

        # Cargar e insertar imágenes de firmas si existen
        autorizacion_path = os.path.join(firmas_folder, row["FirmaAutorizacion"])
        satisfaccion_path = os.path.join(firmas_folder, row["FirmaSatisfaccion"])
        firma_lider_path = os.path.join(firmas_lideres_folder, row["FIRMA_LIDER"])

        if os.path.exists(autorizacion_path):
            replace_text_with_image(doc, "@FIRMA_AUTORIZA@", autorizacion_path)
        else:
            status = "Fallo"
            error_msg += f"Falta imagen de autorización: {autorizacion_path}. "

        if os.path.exists(satisfaccion_path):
            replace_text_with_image(doc, "@FIRMA_SATISFACCION@", satisfaccion_path)
        else:
            status = "Fallo"
            error_msg += f"Falta imagen de satisfacción: {satisfaccion_path}. "
         
        if os.path.exists(firma_lider_path):
            replace_text_with_image(doc, "@FIRMA_LIDER@", firma_lider_path, width=2)
        else:
            status = "Fallo"
            error_msg += f"Falta imagen de satisfacción: {firma_lider_path}. "

        # Guardar documento Word
        doc.save(word_output_path)
        word_files.append(word_output_path)
        pdf_files.append(pdf_output_path)
        
        #Barra de progreso
        print_progress_bar(cuenta_word, total, start_time=start_time)
        cuenta_word += 1

    except Exception as e:
        status = "Fallo"
        error_msg += f"Error en generación Word: {str(e)}. "

    # Guardar información en la lista de resultados
    row_data = row.to_dict()
    row_data["Estado"] = status
    row_data["Errores"] = error_msg
    resultados.append(row_data)

# Crear DataFrame de resultados y guardarlo en Excel
df_resultados = pd.DataFrame(resultados)
df_resultados.to_excel(output_excel_path, index=False)

# Conversión a PDF con barra de progreso
total = len(word_files)-1
start_time = time.time()  # Captura el tiempo de inicio
cuenta_pdf = 0
print("\nProceso de generación de PDFs")
try:
    word_app = comtypes.client.CreateObject('Word.Application')
    word_app.Visible = False  

    for word_path, pdf_path in zip(word_files, pdf_files):
        try:
            doc = word_app.Documents.Open(os.path.abspath(word_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # Guardar como PDF
            doc.Close()
        except Exception as e:
            resultados[cuenta_pdf]["Estado"] = "Fallo"
            resultados[cuenta_pdf]["Errores"] += f"Error en conversión a PDF: {str(e)}. "
        
        if borrar_docx:
            os.remove(word_path)
        print_progress_bar(cuenta_pdf, total, start_time=start_time)
        cuenta_pdf += 1

    word_app.Quit()
except Exception as e:
    print(f"\n❌ Error en la inicialización de Word: {str(e)}")

# Guardar nuevamente el Excel actualizado con errores de PDF
df_resultados = pd.DataFrame(resultados)
df_resultados.to_excel(output_excel_path, index=False)

print(f"\n✅ Proceso completado. Se generaron los documentos y el reporte en Excel de {cuenta_pdf} actas.")

